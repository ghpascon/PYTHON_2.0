import yfinance as yf
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
import numpy as np
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime, timedelta

def get_ticker_data(ticker, start_date, end_date):
    """
    Obtém os dados históricos de um ticker entre as datas fornecidas.
    """
    stock = yf.Ticker(ticker)
    data = stock.history(start=start_date, end=end_date)
    return data

def plot_ticker(ticker, start_date, end_date, file_path):
    """
    Plota os valores de fechamento e a linha de regressão linear para um ticker e salva o gráfico em um arquivo.
    """
    # Obter os dados do ticker
    stock_data = get_ticker_data(ticker, start_date, end_date)
    stock_values = stock_data['Close']
    
    # Convertendo as datas para números inteiros para o cálculo da regressão
    x = np.array(range(len(stock_values))).reshape(-1, 1)  # Datas como números
    y = stock_values.values  # Valores de fechamento
    
    # Criando o modelo de regressão linear
    model = LinearRegression()
    model.fit(x, y)
    
    # Calculando a linha de regressão
    y_pred = model.predict(x)
    
    # Plotando os valores reais de fechamento
    plt.figure(figsize=(10, 6))
    plt.plot(stock_values.index, stock_values.values, label=f"{ticker} Fechamento", color="blue")
    
    # Plotando a linha de regressão
    plt.plot(stock_values.index, y_pred, label=f"{ticker} Regressão Linear", color="red", linestyle="--")
    
    # Adicionar título e rótulos
    plt.title(f"{ticker} - Fechamento com Regressão Linear")
    plt.xlabel("Data")
    plt.ylabel("Fechamento")
    plt.grid(True)
    plt.legend()

    # Salvar o gráfico em um arquivo
    plt.tight_layout()
    plt.savefig(file_path)
    plt.close()

def generate_summary_data(ticker, start_date, end_date):
    """
    Gera as estatísticas de resumo (data inicial, data final, valor na data inicial, valor na data final,
    menor valor, maior valor e porcentagem de aumento) para um ticker entre as datas fornecidas.
    """
    stock_data = get_ticker_data(ticker, start_date, end_date)
    
    # Obter as datas inicial e final
    start_date = stock_data.index[0]  # Primeira data
    end_date = stock_data.index[-1]  # Última data
    
    start_value = stock_data['Close'].iloc[0]  # Valor na data inicial
    end_value = stock_data['Close'].iloc[-1]  # Valor na data final
    
    min_value = stock_data['Close'].min()  # Menor valor
    max_value = stock_data['Close'].max()  # Maior valor
    
    # Cálculo da porcentagem de aumento
    percentage_increase = ((end_value - start_value) / start_value) * 100
    
    return start_date, end_date, start_value, end_value, min_value, max_value, percentage_increase

def create_docx_with_charts(tickers, intervals, docx_filename):
    """
    Cria um arquivo DOCX com gráficos e uma tabela contendo estatísticas para múltiplos tickers.
    """
    # Criar um novo documento
    doc = Document()
    doc.add_heading('Análise de Fechamento com Regressão Linear', 0)
    
    # Gerar gráficos e adicionar ao documento para cada intervalo de tempo
    for ticker in tickers:
        for interval in intervals:
            # Calcular as datas de início e fim com base no intervalo
            end_date = datetime.today()
            start_date = end_date - timedelta(days=interval)
            
            # Caminho do arquivo de imagem
            image_path = f"{ticker}_{interval}_chart.png"
            
            # Gerar o gráfico e salvar em um arquivo
            plot_ticker(ticker, start_date, end_date, image_path)
            
            # Adicionar título para o gráfico
            doc.add_heading(f'Gráfico de {ticker} - Últimos {interval} dias', level=1)
            
            # Inserir a imagem no documento
            doc.add_picture(image_path, width=Inches(5.0))
            
            # Adicionar tabela com as estatísticas do ticker
            start_date, end_date, start_value, end_value, min_value, max_value, percentage_increase = generate_summary_data(ticker, start_date, end_date)
            
            # Criar a tabela para o ticker
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            
            # Adicionar cabeçalho à tabela
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Data Inicial'
            hdr_cells[1].text = 'Valor na Data Inicial'
            hdr_cells[2].text = 'Data Final'
            hdr_cells[3].text = 'Valor na Data Final'
            hdr_cells[4].text = 'Menor Valor'
            hdr_cells[5].text = 'Maior Valor'
            hdr_cells[6].text = 'Aumento (%)'
            
            # Adicionar uma linha com as estatísticas
            row_cells = table.add_row().cells
            row_cells[0].text = str(start_date.date())  # Exibir apenas a data (sem hora)
            row_cells[1].text = f"{start_value:.2f}"  # Valor na data inicial
            row_cells[2].text = str(end_date.date())  # Exibir apenas a data (sem hora)
            row_cells[3].text = f"{end_value:.2f}"  # Valor na data final
            row_cells[4].text = f"{min_value:.2f}"
            row_cells[5].text = f"{max_value:.2f}"
            row_cells[6].text = f"{percentage_increase:.2f}%"
            
            # Adicionar espaçamento
            doc.add_paragraph("\n")
            
            # Remover a imagem temporária
            os.remove(image_path)
    
    # Salvar o documento
    doc.save(docx_filename)
    print(f"Documento {docx_filename} foi criado com sucesso!")

if __name__ == "__main__":
    # Lista de tickers para análise
    tickers = ["^BVSP", "PETR4.SA", "GOAU4.SA", "MSFT"]  # Exemplos de tickers (Ibovespa, Apple, Google, Microsoft)
    
    # Lista de intervalos (em dias)
    intervals = [7, 30, 90, 365]  # Últimos 7, 30, 90 e 365 dias
    
    # Nome do arquivo DOCX
    docx_filename = "analise_tickers_com_intervalos.docx"
    
    # Criar o arquivo DOCX com os gráficos e a tabela para cada intervalo
    create_docx_with_charts(tickers, intervals, docx_filename)
